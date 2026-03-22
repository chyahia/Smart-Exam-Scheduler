from flask import Blueprint, request, jsonify, send_file
import json
import io
from datetime import datetime
from collections import defaultdict
from app.database import get_db_connection

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import pandas as pd
from app.services.algorithms import _run_initial_subject_placement, clean_string_for_matching
from openpyxl.utils import get_column_letter
from openpyxl.styles import Alignment

export_bp = Blueprint('export', __name__)

def create_word_document_with_table(doc, title, headers, data_grid):
    heading = doc.add_heading(level=2)
    heading.clear() 
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    pPr = heading._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)
    
    run = heading.add_run(title)
    font = run.font
    font.rtl = True
    font.name = 'Arial'

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.autofit = False

    tbl_pr = table._element.xpath('w:tblPr')[0]
    bidi_visual_element = OxmlElement('w:bidiVisual')
    tbl_pr.append(bidi_visual_element)

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell_paragraph = hdr_cells[i].paragraphs[0]
        cell_paragraph.text = ""
        run = cell_paragraph.add_run(header)
        font = run.font
        font.rtl = True
        font.name = 'Arial'
        cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_paragraph.paragraph_format.rtl = True

    for row_data in data_grid:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            cell_paragraph = row_cells[i].paragraphs[0]
            cell_paragraph.text = ""
            lines = str(cell_data).split('\n')
            for idx, line in enumerate(lines):
                if idx > 0:
                    cell_paragraph.add_run().add_break()
                run = cell_paragraph.add_run(line)
                font = run.font
                font.rtl = True
                font.name = 'Arial'
            cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cell_paragraph.paragraph_format.rtl = True
            
    doc.add_page_break()

@export_bp.route('/api/export/word/all-exams', methods=['POST'])
def export_exams_word():
    schedule_data = request.get_json()
    if not schedule_data: return jsonify({"error": "No schedule data provided"}), 400
    
    conn = get_db_connection()
    # تحديث اسم الجدول إلى professor_subject
    assignments_rows = conn.execute('''
        SELECT s.name as subj_name, l.name as level_name, p.name as prof_name 
        FROM professor_subject ps 
        JOIN subjects s ON ps.subject_id = s.id 
        JOIN levels l ON s.level_id = l.id 
        JOIN professors p ON ps.professor_id = p.id
    ''').fetchall()
    settings_row = conn.execute("SELECT value FROM settings WHERE key = 'main_settings'").fetchone()
    conn.close()
    
    settings_data = json.loads(settings_row['value']) if settings_row else {}
    guards_large = int(settings_data.get('guardsLargeHall', 4))
    guards_medium = int(settings_data.get('guardsMediumHall', 2))
    guards_small = int(settings_data.get('guardsSmallHall', 1))

    subject_owners = {(row['subj_name'], row['level_name']): row['prof_name'] for row in assignments_rows}
    
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    margin = Cm(0.5)
    section.top_margin, section.bottom_margin, section.left_margin, section.right_margin = margin, margin, margin, margin

    all_dates = sorted(schedule_data.keys())
    all_times = sorted({time for date_slots in schedule_data.values() for time in date_slots})
    all_levels = sorted({exam['level'] for slots in schedule_data.values() for exams in slots.values() for exam in exams})
    day_names = ["الأحد", "الاثنين", "الثلاثاء", "الأربعاء", "الخميس", "الجمعة", "السبت"]
    
    headers = ["الفترة"] + [f"{day_names[datetime.strptime(d, '%Y-%m-%d').isoweekday() % 7]}\n{d}" for d in all_dates]

    for level in all_levels:
        data_grid = []
        for time in all_times:
            row_data = [time]
            for date in all_dates:
                exam = next((e for e in schedule_data.get(date, {}).get(time, []) if e['level'] == level), None)
                content = ""
                if exam:
                    owner = subject_owners.get((exam['subject'], exam['level']), "غير محدد")
                    content = f"{exam['subject']}\nأستاذ المادة: {owner}\n\nالحراسة:"

                    halls_by_type = defaultdict(list)
                    for h in exam.get('halls', []): halls_by_type[h['type']].append(h['name'])
                    
                    guards_copy = [g for g in exam.get('guards', []) if g != "**نقص**"]

                    if halls_by_type.get('كبيرة'):
                        num_guards_needed = len(halls_by_type['كبيرة']) * guards_large
                        g_list = guards_copy[:num_guards_needed]
                        guards_copy = guards_copy[num_guards_needed:]
                        hall_names = ", ".join(halls_by_type['كبيرة'])
                        guard_text = '\n'.join(g_list) if g_list else '(لا يوجد)'
                        content += f"\nالقاعة الكبيرة: {hall_names}\n{guard_text}"
                    
                    other_hall_names = halls_by_type.get('متوسطة', []) + halls_by_type.get('صغيرة', [])
                    if other_hall_names:
                        guard_text = '\n'.join(guards_copy) if guards_copy else '(لا يوجد)'
                        content += f"\nالقاعات الأخرى: {', '.join(other_hall_names)}\n{guard_text}"
                
                row_data.append(content)
            data_grid.append(row_data)
        
        create_word_document_with_table(doc, f"جدول امتحانات: {level}", headers, data_grid)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="جداول_الامتحانات.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@export_bp.route('/api/export/word/all-profs', methods=['POST'])
def export_profs_word():
    schedule_data = request.get_json()
    if not schedule_data: return jsonify({"error": "No schedule data provided"}), 400

    conn = get_db_connection()
    all_professors = sorted([p['name'] for p in conn.execute("SELECT name FROM professors").fetchall()])
    assignments_rows = conn.execute('''
        SELECT p.name as prof_name, s.name as subj_name, l.name as level_name 
        FROM professor_subject ps 
        JOIN professors p ON ps.professor_id = p.id 
        JOIN subjects s ON ps.subject_id = s.id 
        JOIN levels l ON s.level_id = l.id
    ''').fetchall()
    conn.close()

    prof_owned_subjects = defaultdict(set)
    for row in assignments_rows:
        prof_owned_subjects[row['prof_name']].add((row['subj_name'], row['level_name']))

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    margin = Cm(0.5)
    section.top_margin, section.bottom_margin, section.left_margin, section.right_margin = margin, margin, margin, margin

    all_dates = sorted(schedule_data.keys())
    all_times = sorted({time for date_slots in schedule_data.values() for time in date_slots})
    day_names = ["الأحد", "الاثنين", "الثلاثاء", "الأربعاء", "الخميس", "الجمعة", "السبت"]
    
    for prof_name in all_professors:
        title = f"جدول الحراسة: {prof_name}"
        headers = ["اليوم/التاريخ"] + all_times
        
        heading = doc.add_heading(level=2); heading.clear()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr = heading._p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi'); bidi.set(qn('w:val'), '1'); pPr.append(bidi)
        run = heading.add_run(title)
        font = run.font; font.rtl = True; font.name = 'Arial'

        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        table.autofit = False
        tbl_pr = table._element.xpath('w:tblPr')[0]
        bidi_visual_element = OxmlElement('w:bidiVisual')
        tbl_pr.append(bidi_visual_element)
        
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            p = hdr_cells[i].paragraphs[0]; p.text = ""
            run = p.add_run(header)
            font = run.font; font.rtl = True; font.name = 'Arial'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.rtl = True

        has_any_duty = False
        for date in all_dates:
            row_cells = table.add_row().cells
            day_name = day_names[datetime.strptime(date, '%Y-%m-%d').isoweekday() % 7]
            
            p = row_cells[0].paragraphs[0]; p.text = ""
            run = p.add_run(f"{day_name}\n{date}"); run.font.rtl = True; run.font.name = 'Arial'; run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.rtl = True

            for i, time in enumerate(all_times, 1):
                cell_content_parts = []
                is_teaching_and_guarding = False
                is_teaching_only = False
                
                exams_in_slot = schedule_data.get(date, {}).get(time, [])
                
                for exam in exams_in_slot:
                    is_guarding = prof_name in exam.get('guards', [])
                    is_owner = (exam['subject'], exam['level']) in prof_owned_subjects.get(prof_name, set())

                    if is_guarding or is_owner:
                        has_any_duty = True
                        if is_guarding:
                            if is_owner: is_teaching_and_guarding = True
                            cell_content_parts.append(f"{exam['subject']} ({exam['level']})\n(حراسة)")
                        elif is_owner:
                            is_teaching_only = True
                            cell_content_parts.append(f"{exam['subject']} ({exam['level']})\n(دون حراسة)")
                
                p = row_cells[i].paragraphs[0]; p.text = ""
                lines = "\n---\n".join(cell_content_parts).split('\n')
                for idx, line in enumerate(lines):
                    if idx > 0: p.add_run().add_break()
                    run = p.add_run(line)
                    font = run.font; font.rtl = True; font.name = 'Arial'
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; p.paragraph_format.rtl = True
                
                shading_elm = OxmlElement('w:shd')
                if is_teaching_and_guarding:
                    shading_elm.set(qn('w:fill'), 'D4EDDA') # تظليل أخضر خفيف
                    row_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
                elif is_teaching_only:
                    shading_elm.set(qn('w:fill'), 'FFF3CD') # تظليل أصفر خفيف
                    row_cells[i]._tc.get_or_add_tcPr().append(shading_elm)

        if has_any_duty:
             doc.add_page_break()
        else:
            doc._body.remove(table._element)
            doc._body.remove(heading._element)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="جداول_الحراسة_للأساتذة.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@export_bp.route('/api/export/word/all-profs-anonymous', methods=['POST'])
def export_profs_anonymous_word():
    schedule_data = request.get_json()
    if not schedule_data: return jsonify({"error": "No schedule data provided"}), 400

    conn = get_db_connection()
    all_professors = sorted([p['name'] for p in conn.execute("SELECT name FROM professors").fetchall()])
    assignments_rows = conn.execute('''
        SELECT p.name as prof_name, s.name as subj_name, l.name as level_name 
        FROM professor_subject ps 
        JOIN professors p ON ps.professor_id = p.id 
        JOIN subjects s ON ps.subject_id = s.id 
        JOIN levels l ON s.level_id = l.id
    ''').fetchall()
    conn.close()

    prof_owned_subjects = defaultdict(set)
    for row in assignments_rows:
        prof_owned_subjects[row['prof_name']].add((row['subj_name'], row['level_name']))

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    margin = Cm(0.5)
    section.top_margin, section.bottom_margin, section.left_margin, section.right_margin = margin, margin, margin, margin

    all_dates = sorted(schedule_data.keys())
    all_times = sorted({time for date_slots in schedule_data.values() for time in date_slots})
    day_names = ["الأحد", "الاثنين", "الثلاثاء", "الأربعاء", "الخميس", "الجمعة", "السبت"]
    
    for prof_name in all_professors:
        title = f"جدول الحراسة (مُبسَّط): {prof_name}"
        headers = ["اليوم/التاريخ"] + all_times
        
        heading = doc.add_heading(level=2); heading.clear()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr = heading._p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi'); bidi.set(qn('w:val'), '1'); pPr.append(bidi)
        run = heading.add_run(title)
        font = run.font; font.rtl = True; font.name = 'Arial'

        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        table.autofit = False
        tbl_pr = table._element.xpath('w:tblPr')[0]
        bidi_visual_element = OxmlElement('w:bidiVisual')
        tbl_pr.append(bidi_visual_element)
        
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            p = hdr_cells[i].paragraphs[0]; p.text = ""
            run = p.add_run(header)
            font = run.font; font.rtl = True; font.name = 'Arial'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.rtl = True

        has_any_duty = False
        for date in all_dates:
            row_cells = table.add_row().cells
            day_name = day_names[datetime.strptime(date, '%Y-%m-%d').isoweekday() % 7]
            
            p = row_cells[0].paragraphs[0]; p.text = ""
            run = p.add_run(f"{day_name}\n{date}"); run.font.rtl = True; run.font.name = 'Arial'; run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.rtl = True

            for i, time in enumerate(all_times, 1):
                cell_content_parts = []
                is_teaching_and_guarding = False
                is_teaching_only = False
                
                exams_in_slot = schedule_data.get(date, {}).get(time, [])
                
                for exam in exams_in_slot:
                    is_guarding = prof_name in exam.get('guards', [])
                    is_owner = (exam['subject'], exam['level']) in prof_owned_subjects.get(prof_name, set())

                    if is_guarding or is_owner:
                        has_any_duty = True
                        if is_guarding:
                            if is_owner:
                                is_teaching_and_guarding = True
                                cell_content_parts.append(f"{exam['subject']} ({exam['level']})\n(حراسة)")
                            else:
                                cell_content_parts.append("(تكليف بحراسة)")
                        elif is_owner:
                            is_teaching_only = True
                            cell_content_parts.append(f"{exam['subject']} ({exam['level']})\n(دون حراسة)")
                
                p = row_cells[i].paragraphs[0]; p.text = ""
                lines = "\n---\n".join(cell_content_parts).split('\n')
                for idx, line in enumerate(lines):
                    if idx > 0: p.add_run().add_break()
                    run = p.add_run(line)
                    font = run.font; font.rtl = True; font.name = 'Arial'
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; p.paragraph_format.rtl = True
                
                shading_elm = OxmlElement('w:shd')
                if is_teaching_and_guarding:
                    shading_elm.set(qn('w:fill'), 'D4EDDA')
                    row_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
                elif is_teaching_only:
                    shading_elm.set(qn('w:fill'), 'FFF3CD')
                    row_cells[i]._tc.get_or_add_tcPr().append(shading_elm)

        if has_any_duty:
             doc.add_page_break()
        else:
            doc._body.remove(table._element)
            doc._body.remove(heading._element)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="جداول_الحراسة_المبسطة.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@export_bp.route('/api/export-manual-distribution-template', methods=['POST'])
def export_manual_distribution_template():
    try:
        conn = get_db_connection()
        
        # 1. جلب جدول الأيام والفترات مباشرة من قاعدة البيانات
        row_sched = conn.execute("SELECT value FROM settings WHERE key = 'exam_schedule'").fetchone()
        exam_schedule = json.loads(row_sched['value']) if row_sched else {}
        
        all_dates = sorted(exam_schedule.keys())
        all_times = sorted(list(set(time for slots in exam_schedule.values() for slot in slots for time in [slot['time']])))
        
        if not all_dates or not all_times:
            conn.close()
            return jsonify({"error": "الرجاء حفظ جدول الأيام والفترات في المرحلة 4 أولاً لتتمكن من تصدير المخطط."}), 400

        # 2. جلب باقي البيانات
        all_levels_list = [row['name'] for row in conn.execute("SELECT name FROM levels").fetchall()]
        all_subjects_rows = conn.execute("SELECT s.name as subj_name, l.name as level_name FROM subjects s JOIN levels l ON s.level_id = l.id").fetchall()
        all_subjects = [{'name': row['subj_name'], 'level': row['level_name']} for row in all_subjects_rows]
        all_halls = [dict(row) for row in conn.execute("SELECT name, type FROM halls").fetchall()]
        
        assignments_rows = conn.execute('''
            SELECT p.name as prof_name, s.name as subj_name, l.name as level_name 
            FROM professor_subject ps 
            JOIN professors p ON ps.professor_id = p.id 
            JOIN subjects s ON ps.subject_id = s.id 
            JOIN levels l ON s.level_id = l.id
        ''').fetchall()

        lh_rows = conn.execute('''
            SELECT l.name as level_name, h.name as hall_name
            FROM level_halls lh 
            JOIN levels l ON lh.level_id = l.id 
            JOIN halls h ON lh.hall_id = h.id
        ''').fetchall()
        conn.close()
        
        subject_owners = { (clean_string_for_matching(s['subj_name']), clean_string_for_matching(s['level_name'])): clean_string_for_matching(s['prof_name']) for s in assignments_rows }

        level_hall_assignments = defaultdict(list)
        for row in lh_rows:
            level_hall_assignments[row['level_name']].append(row['hall_name'])

        settings_for_placement = {
            'examSchedule': exam_schedule,
            'levelHallAssignments': dict(level_hall_assignments)
        }

        # 3. تشغيل خوارزمية التوزيع الأولي
        initial_schedule, _ = _run_initial_subject_placement(settings_for_placement, all_subjects, all_levels_list, subject_owners, all_halls)

        # 4. بناء ملف الإكسل
        output = io.BytesIO()
        writer = pd.ExcelWriter(output, engine='openpyxl')
        
        for level_name in sorted(all_levels_list):
            df_level = pd.DataFrame(index=all_times, columns=all_dates)
            df_level.index.name = "الفترة"
            
            for date, slots in initial_schedule.items():
                for time, exams in slots.items():
                    for exam in exams:
                        if exam['level'] == level_name:
                            cell_content = f"{exam['subject']} ::: {exam['professor']} ::: {exam['level']}"
                            df_level.at[time, date] = cell_content
            
            unplaced_subjects = [s for s in all_subjects if s['level'] == level_name and not any(e['subject'] == s['name'] and e['level'] == s['level'] for d in initial_schedule.values() for t in d.values() for e in t)]
            if unplaced_subjects:
                unplaced_row_name = "--- مواد غير موزعة ---"
                df_level.loc[unplaced_row_name] = ''
                cell_texts = [f"{s['name']} ::: {subject_owners.get((s['name'], s['level']), 'غير محدد')} ::: {s['level']}" for s in unplaced_subjects]
                if all_dates:
                    df_level.at[unplaced_row_name, all_dates[0]] = "\n".join(cell_texts)

            safe_sheet_name = level_name[:31]
            df_level.to_excel(writer, sheet_name=safe_sheet_name)
            worksheet = writer.sheets[safe_sheet_name]

            worksheet.sheet_view.rightToLeft = True
            
            # 1. عرض العمود الأول (الفترات الزمنية)
            worksheet.column_dimensions['A'].width = 18
            
            # 2. تقليل عرض أعمدة الأيام (قللناه من 35 إلى 25، يمكنك تغييره كما تشاء)
            for i in range(2, len(all_dates) + 2):
                worksheet.column_dimensions[get_column_letter(i)].width = 18
                
            # جعلنا النص يتوسط الخلية عمودياً ليكون شكله أجمل مع الارتفاع الجديد
            wrap_alignment = Alignment(wrap_text=True, horizontal='right', vertical='center')
            
            for row in worksheet.iter_rows():
                # 3. زيادة ارتفاع كل صف (وضعناه 60، ويمكنك زيادته إلى 70 أو 80 حسب رغبتك)
                worksheet.row_dimensions[row[0].row].height = 80
                
                for cell in row:
                    cell.alignment = wrap_alignment

        writer.close()
        
        excel_data = output.getvalue()
        final_output = io.BytesIO(excel_data)
        
        return send_file(
            final_output, 
            as_attachment=True, 
            download_name='مخطط_توزيع_المواد_للتعديل.xlsx', 
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500
    
@export_bp.route('/api/import-manual-distribution', methods=['POST'])
def import_manual_distribution():
    if 'file' not in request.files: return jsonify({"error": "لم يتم العثور على ملف."}), 400
    file = request.files['file']
    try:
        xls = pd.read_excel(file, sheet_name=None, index_col=0, dtype=str)
        pinned_schedule = defaultdict(lambda: defaultdict(list))
        
        conn = get_db_connection()
        all_halls = [dict(row) for row in conn.execute("SELECT id, name, type FROM halls").fetchall()]
        
        # استخراج قاعات المستويات لربطها
        lh_rows = conn.execute('''
            SELECT l.name as level_name, h.id as hall_id, h.name as hall_name, h.type as hall_type
            FROM level_halls lh 
            JOIN levels l ON lh.level_id = l.id 
            JOIN halls h ON lh.hall_id = h.id
        ''').fetchall()
        conn.close()

        level_hall_assignments = defaultdict(list)
        for lh in lh_rows:
            level_hall_assignments[lh['level_name']].append({'name': lh['hall_name'], 'type': lh['hall_type']})

        pinned_count = 0
        for sheet_name, df in xls.items():
            for date in df.columns:
                for time in df.index:
                    cell_value = df.at[time, date]
                    if pd.notna(cell_value):
                        subjects_in_cell = cell_value.strip().split('\n')
                        for subject_line in subjects_in_cell:
                            if ':::' in subject_line:
                                try:
                                    subject_name, professor_name, level_name = [part.strip() for part in subject_line.split(' ::: ')]
                                    if not date or not time or "مواد غير موزعة" in time: continue

                                    halls_details = level_hall_assignments.get(level_name, [])
                                    exam = {
                                        "date": date.strip(), "time": time.strip(),
                                        "subject": subject_name, "level": level_name,
                                        "professor": professor_name, "halls": halls_details,
                                        "guards": []
                                    }
                                    pinned_schedule[exam['date']][exam['time']].append(exam)
                                    pinned_count += 1
                                except ValueError:
                                    continue
        
        conn = get_db_connection()
        conn.execute("INSERT OR REPLACE INTO settings (key, value) VALUES (?, ?)", 
                     ('pinned_subject_schedule', json.dumps(pinned_schedule)))
        conn.commit()
        conn.close()
        return jsonify({"success": True, "message": f"تم استيراد وتثبيت {pinned_count} مادة بنجاح."})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@export_bp.route('/api/clear-manual-distribution', methods=['POST'])
def clear_manual_distribution():
    try:
        conn = get_db_connection()
        conn.execute("DELETE FROM settings WHERE key = 'pinned_subject_schedule'")
        conn.commit()
        conn.close()
        return jsonify({"success": True, "message": "تم مسح الجدول اليدوي. سيعتمد التشغيل القادم على التوزيع التلقائي."})
    except Exception as e:
        return jsonify({"error": str(e)}), 500